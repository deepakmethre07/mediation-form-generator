from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ---------- TITLE ----------
def add_centered_bold(text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_centered_bold("FORM ‘A’", 12)
add_centered_bold("MEDIATION APPLICATION FORM", 12)
add_centered_bold("[REFER RULE 3(1)]", 10)
doc.add_paragraph("Mumbai District Legal Services Authority").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("City Civil Court, Mumbai").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\nDETAILS OF PARTIES:")

# ---------- TABLE ----------
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Sr. No."
hdr_cells[1].text = "Particulars"
hdr_cells[2].text = "Details"

def add_row(c1, c2, c3):
    row = table.add_row().cells
    row[0].text = c1
    row[1].text = c2
    row[2].text = c3

add_row("1", "Name of Applicant", "{{client_name}}")
add_row("", "Address and contact details of Applicant", "")
add_row("", "REGISTERED ADDRESS", "{{branch_address}}")
add_row("", "CORRESPONDENCE BRANCH ADDRESS", "{{branch_address}}")
add_row("", "Telephone No.", "{{mobile}}")
add_row("", "Mobile No.", "")
add_row("", "Email ID", "info@kslegal.co.in")

add_row("2", "Name of Opposite Party", "{{customer_name}}")
add_row("", "REGISTERED ADDRESS", "{{address1 or '______________'}}")
add_row("", "CORRESPONDENCE ADDRESS", "{{address1 or '______________'}}")
add_row("", "Telephone No.", "")
add_row("", "Mobile No.", "")
add_row("", "Email ID", "")

doc.add_paragraph("\nDETAILS OF DISPUTE:")

p = doc.add_paragraph()
run = p.add_run("THE COMM. COURTS (PRE-INSTITUTION SETTLEMENT) RULES, 2018")
run.bold = True

doc.add_paragraph(
    "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):"
)

# ---------- SAVE ----------
doc.save("FORM_A_Mediation_Application.docx")
print("FORM A Word document created successfully.")
